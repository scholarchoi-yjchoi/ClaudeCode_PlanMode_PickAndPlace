#!/usr/bin/env python3
"""
Isaac Sim 4.5 Pick and Place Project - Technical Report Generator
Generates a comprehensive technical documentation in DOCX format.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_with_style(doc, text, level):
    """Add heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    return heading


def create_document():
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("\n\n\n\nIsaac Sim 4.5\nPick and Place Project")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Technical Report")
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = """
Franka Panda Robot + YCB Dataset + RMPFlow Controller

Version 3.0
January 2026

Project Duration: 2026.01.14 ~ 2026.01.19
Tests Conducted: 26
Final Accuracy: 7.1mm (DynamicCuboid) / 11.1mm (YCB Objects)
"""
    info_run = info.add_run(info_text)
    info_run.font.size = Pt(12)
    info_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    add_heading_with_style(doc, "Table of Contents", 1)

    toc_items = [
        ("1. Executive Summary", 3),
        ("2. Project Overview", 4),
        ("   2.1 Objectives", 4),
        ("   2.2 Key Achievements", 4),
        ("   2.3 Technology Stack", 4),
        ("3. System Requirements", 5),
        ("4. Project Architecture", 6),
        ("   4.1 Directory Structure", 6),
        ("   4.2 Main Components", 6),
        ("5. Core Concepts", 7),
        ("   5.1 Pick and Place Operation", 7),
        ("   5.2 10 Phase Workflow", 7),
        ("   5.3 RMPFlow Controller", 8),
        ("   5.4 Coordinate System", 8),
        ("6. Problem Analysis & Solutions", 9),
        ("   6.1 Problem 1: Wrong Robot Position", 9),
        ("   6.2 Problem 2: Z Height Too High", 9),
        ("   6.3 Problem 3: No Physics on Object", 10),
        ("   6.4 Problem 4: Position Mismatch After Settle", 10),
        ("   6.5 Problem 5: YCB Object Rolling", 11),
        ("   6.6 Problem 6: RMPFlow Table Collision", 11),
        ("   6.7 Problem 7: Gripper Initialization (Critical)", 12),
        ("   6.8 Problem 8: YCB Collision Avoidance", 12),
        ("7. YCB Dataset Support", 13),
        ("   7.1 Collision Proxy Pattern", 13),
        ("   7.2 High-Friction Material", 14),
        ("   7.3 Supported Objects", 14),
        ("8. API Reference", 15),
        ("   8.1 Key Imports", 15),
        ("   8.2 VerificationTracker Class", 15),
        ("   8.3 Configuration Constants", 16),
        ("9. Test Results", 17),
        ("10. How to Run", 18),
        ("11. Future Improvements", 19),
        ("12. References", 19),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run("\t" * 8 + str(page))

    doc.add_page_break()

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================
    add_heading_with_style(doc, "1. Executive Summary", 1)

    doc.add_paragraph(
        "This technical report documents the development of a Pick and Place demonstration "
        "using NVIDIA Isaac Sim 4.5 and the Franka Panda robot manipulator. The project "
        "implements a complete robotic manipulation pipeline that successfully picks objects "
        "from one location and places them in a target container."
    )

    doc.add_paragraph(
        "Key Results:", style='List Bullet'
    )

    results = [
        "Successfully completed Pick and Place operation after 26 iterative tests",
        "Achieved 7.1mm accuracy with DynamicCuboid objects",
        "Extended support to YCB dataset objects with 11.1mm accuracy",
        "Solved 8 critical technical problems during development",
        "Developed the 'Collision Proxy Pattern' for YCB object manipulation"
    ]

    for result in results:
        doc.add_paragraph(result, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        "The most significant discovery was the 'Collision Proxy Pattern' - a technique that "
        "separates visual representation from physics collision, enabling the RMPFlow controller "
        "to accurately approach complex YCB objects without excessive collision avoidance."
    )

    doc.add_page_break()

    # =========================================================================
    # 2. PROJECT OVERVIEW
    # =========================================================================
    add_heading_with_style(doc, "2. Project Overview", 1)

    doc.add_heading("2.1 Objectives", 2)
    doc.add_paragraph(
        "The primary objective of this project is to implement a robust Pick and Place "
        "demonstration using Isaac Sim 4.5's new API (isaacsim.core.api.*). Specific goals include:"
    )

    objectives = [
        "Demonstrate Franka Panda robot picking objects from a table and placing them in a container",
        "Implement real-time verification of robot actions (success/failure detection)",
        "Support multiple object types including YCB dataset objects",
        "Achieve millimeter-level placement accuracy"
    ]

    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading("2.2 Key Achievements", 2)

    # Achievement table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    achievements = [
        ("Metric", "Result"),
        ("Total Tests to Success", "26"),
        ("DynamicCuboid Accuracy", "7.1mm"),
        ("YCB Objects Accuracy", "11.1mm"),
        ("Problems Solved", "8"),
    ]

    for i, (col1, col2) in enumerate(achievements):
        row = table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_heading("2.3 Technology Stack", 2)

    tech_table = doc.add_table(rows=6, cols=2)
    tech_table.style = 'Table Grid'

    tech_stack = [
        ("Component", "Technology"),
        ("Simulation Platform", "NVIDIA Isaac Sim 4.5"),
        ("Robot", "Franka Panda (7-DOF Manipulator)"),
        ("Motion Planner", "RMPFlow Controller"),
        ("Objects", "DynamicCuboid, YCB Dataset"),
        ("API", "isaacsim.core.api.* (New API)"),
    ]

    for i, (col1, col2) in enumerate(tech_stack):
        row = tech_table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # =========================================================================
    # 3. SYSTEM REQUIREMENTS
    # =========================================================================
    add_heading_with_style(doc, "3. System Requirements", 1)

    doc.add_heading("Hardware Requirements", 2)
    hw_reqs = [
        "NVIDIA RTX GPU (RTX 3070 or higher recommended)",
        "32GB RAM minimum",
        "100GB SSD storage",
        "Ubuntu 22.04 LTS"
    ]
    for req in hw_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading("Software Requirements", 2)
    sw_reqs = [
        "NVIDIA Isaac Sim 4.5",
        "Python 3.10+",
        "NVIDIA Driver 535+",
        "USD (Universal Scene Description) runtime"
    ]
    for req in sw_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading("Python Dependencies", 2)
    doc.add_paragraph(
        "The following packages are included with Isaac Sim's Python environment:"
    )
    deps = [
        "numpy - Numerical computations",
        "pxr (USD) - Scene description and manipulation",
        "isaacsim.core.api - Isaac Sim core functionality",
        "isaacsim.robot.manipulators - Robot manipulator support"
    ]
    for dep in deps:
        doc.add_paragraph(dep, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 4. PROJECT ARCHITECTURE
    # =========================================================================
    add_heading_with_style(doc, "4. Project Architecture", 1)

    doc.add_heading("4.1 Directory Structure", 2)

    structure = """
ClaudeCode_PlanMode_PickAndPlace/
├── franka_ycb_pick_place.py      # Main GUI script
├── test_headless.py              # Headless test script
├── work_log.txt                  # Development log
├── README.md                     # Project documentation
├── .gitignore                    # Git ignore patterns
├── documents/                    # Documentation folder
│   ├── 260114/                   # Jan 14 documents
│   ├── 260115_1/                 # Jan 15 session 1
│   ├── 260115_2/                 # Jan 15 session 2
│   └── 260119/                   # Jan 19 documents
└── project_presentation_v3.html  # 3D presentation
"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(structure)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_heading("4.2 Main Components", 2)

    components_table = doc.add_table(rows=6, cols=2)
    components_table.style = 'Table Grid'

    components = [
        ("File", "Description"),
        ("franka_ycb_pick_place.py", "Main script with GUI support. Implements Pick and Place with VerificationTracker, YCB object support, and keyboard controls."),
        ("test_headless.py", "Automated testing script that runs without GUI. Used for rapid iteration and validation."),
        ("VerificationTracker", "Class that tracks robot actions and verifies success/failure by monitoring object positions."),
        ("SimulationController", "Handles keyboard input, simulation state, and user interaction."),
        ("Collision Proxy System", "Creates invisible DynamicCuboid proxies for YCB objects to enable proper gripper approach."),
    ]

    for i, (col1, col2) in enumerate(components):
        row = components_table.rows[i]
        row.cells[0].text = col1
        row.cells[1].text = col2
        if i == 0:
            set_cell_shading(row.cells[0], "003366")
            set_cell_shading(row.cells[1], "003366")
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # =========================================================================
    # 5. CORE CONCEPTS
    # =========================================================================
    add_heading_with_style(doc, "5. Core Concepts", 1)

    doc.add_heading("5.1 Pick and Place Operation", 2)
    doc.add_paragraph(
        "Pick and Place is a fundamental robotic manipulation task where a robot arm:"
    )
    steps = [
        "Moves to a position above the target object",
        "Lowers the end effector to grasp height",
        "Closes the gripper to secure the object",
        "Lifts and transports the object to the destination",
        "Releases the object at the target location"
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("5.2 10 Phase Workflow", 2)
    doc.add_paragraph(
        "The PickPlaceController divides the operation into 10 distinct phases:"
    )

    phase_table = doc.add_table(rows=11, cols=3)
    phase_table.style = 'Table Grid'

    phases = [
        ("Phase", "Name", "Description"),
        ("0", "Moving above pick", "End effector moves to position above the object"),
        ("1", "Lowering to grasp", "EE descends to grasp height (XY accuracy check)"),
        ("2", "Waiting settle", "Brief pause for physics stabilization"),
        ("3", "Closing gripper", "Gripper fingers close to secure object"),
        ("4", "Lifting object", "EE lifts with grasped object"),
        ("5", "Moving to place XY", "Horizontal movement to place location"),
        ("6", "Lowering to place", "EE descends to release height"),
        ("7", "Opening gripper", "Gripper opens to release object"),
        ("8", "Lifting up", "EE moves up after release"),
        ("9", "Going home", "Robot returns to home configuration"),
    ]

    for i, (p, name, desc) in enumerate(phases):
        row = phase_table.rows[i]
        row.cells[0].text = p
        row.cells[1].text = name
        row.cells[2].text = desc
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_heading("5.3 RMPFlow Controller", 2)
    doc.add_paragraph(
        "RMPFlow (Riemannian Motion Policy Flow) is a motion planning framework that:"
    )
    rmpflow_features = [
        "Generates smooth, collision-aware trajectories in real-time",
        "Combines multiple motion policies (reach target, avoid obstacles, joint limits)",
        "Uses potential fields for obstacle avoidance",
        "Provides reactive control for dynamic environments"
    ]
    for feature in rmpflow_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Important Discovery: ").bold = True
    p.add_run(
        "RMPFlow's collision avoidance can be overly conservative with complex mesh "
        "geometries, causing the gripper to stop 6cm above the target. This led to the "
        "development of the Collision Proxy Pattern."
    )

    doc.add_heading("5.4 Coordinate System", 2)
    doc.add_paragraph("Isaac Sim uses a right-handed coordinate system:")
    coords = [
        "X-axis: Forward (towards robot's front)",
        "Y-axis: Left (from robot's perspective)",
        "Z-axis: Up (vertical)"
    ]
    for coord in coords:
        doc.add_paragraph(coord, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph("Key positions used in this project:")
    pos_table = doc.add_table(rows=4, cols=2)
    pos_table.style = 'Table Grid'
    positions = [
        ("Position", "Coordinates [X, Y, Z]"),
        ("Object spawn", "[0.3, 0.3, 0.3]"),
        ("Place target", "[0.3, -0.3, 0.27]"),
        ("Container", "[0.3, -0.3, 0.02]"),
    ]
    for i, (name, coords) in enumerate(positions):
        row = pos_table.rows[i]
        row.cells[0].text = name
        row.cells[1].text = coords
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # =========================================================================
    # 6. PROBLEM ANALYSIS & SOLUTIONS
    # =========================================================================
    add_heading_with_style(doc, "6. Problem Analysis & Solutions", 1)

    doc.add_paragraph(
        "During development, 8 major technical problems were identified and solved. "
        "This section documents each problem, its root cause, and the solution implemented."
    )

    # Problem 1
    doc.add_heading("6.1 Problem 1: Wrong Robot Position", 2)
    p = doc.add_paragraph()
    p.add_run("Symptom: ").bold = True
    p.add_run("End Effector XY distance was 0.5417m - robot moved to completely wrong location.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Root Cause: ").bold = True
    p.add_run("Object position was outside Franka robot's reachable workspace.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Solution: ").bold = True
    p.add_run("Adjusted TABLE_POSITION to [0.35, 0.0, 0.25] and PLACEMENT_BOUNDS to x:[0.25~0.45], y:[0.10~0.30] based on official example positions (~[0.3, 0.3, 0.3]).")

    # Problem 2
    doc.add_heading("6.2 Problem 2: Z Height Too High", 2)
    p = doc.add_paragraph()
    p.add_run("Symptom: ").bold = True
    p.add_run("EE z=0.36m vs object z=0.27m - robot couldn't reach low enough.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Root Cause: ").bold = True
    p.add_run("Table height was too low for the robot to reach the object at that XY position.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Solution: ").bold = True
    p.add_run("Adjusted table height to z=0.25 and object height to z=0.27, within robot's reachable envelope.")

    # Problem 3
    doc.add_heading("6.3 Problem 3: No Physics on Object", 2)
    p = doc.add_paragraph()
    p.add_run("Symptom: ").bold = True
    p.add_run("Robot completed motion but object didn't move at all.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Root Cause: ").bold = True
    p.add_run("YCB objects created with prim_utils.create_prim() lacked physics attributes.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Solution: ").bold = True
    doc.add_paragraph()

    code = """from pxr import UsdPhysics, PhysxSchema

# Apply physics to YCB object
UsdPhysics.RigidBodyAPI.Apply(prim)
UsdPhysics.CollisionAPI.Apply(prim)
mass_api = UsdPhysics.MassAPI.Apply(prim)
mass_api.CreateMassAttr(0.1)  # 100g"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    # Problem 4
    doc.add_heading("6.4 Problem 4: Position Mismatch After Settle", 2)
    p = doc.add_paragraph()
    p.add_run("Symptom: ").bold = True
    p.add_run("Object moved but to wrong location (0.2728m error).")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Root Cause: ").bold = True
    p.add_run("After physics enabled, object settled to different position, but robot still targeted initial saved position.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Solution: ").bold = True
    doc.add_paragraph()

    code = """# After world.reset(), get actual object position
for i in range(50):  # Physics settle steps
    my_world.step(render=False)

object_position = verifier.get_object_position().copy()
verifier.pick_target = object_position  # Use settled position"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_page_break()

    # Problem 5
    doc.add_heading("6.5 Problem 5: YCB Object Rolling", 2)
    p = doc.add_paragraph()
    p.add_run("Symptom: ").bold = True
    p.add_run("YCB banana rolled off the table during physics settle.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Root Cause: ").bold = True
    p.add_run("Irregular shape of YCB banana caused instability on flat surface.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Solution: ").bold = True
    p.add_run("Switched to DynamicCuboid for stable physics behavior. Later implemented Collision Proxy Pattern for YCB support.")

    # Problem 6
    doc.add_heading("6.6 Problem 6: RMPFlow Table Collision", 2)
    p = doc.add_paragraph()
    p.add_run("Symptom: ").bold = True
    p.add_run("With table present, EE XY distance was 0.45~0.55m (completely wrong position).")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Root Cause: ").bold = True
    p.add_run("RMPFlow controller recognized table as obstacle and generated avoidance paths.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Key Discovery: ").bold = True
    doc.add_paragraph()

    discovery_table = doc.add_table(rows=3, cols=2)
    discovery_table.style = 'Table Grid'
    discoveries = [
        ("Condition", "EE XY Distance"),
        ("Without table", "0.0147m (Accurate!)"),
        ("With table", "0.45~0.55m (Failed)"),
    ]
    for i, (cond, dist) in enumerate(discoveries):
        row = discovery_table.rows[i]
        row.cells[0].text = cond
        row.cells[1].text = dist
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Solution: ").bold = True
    p.add_run("Removed table, used only GroundPlane. Official examples also use GroundPlane without separate table obstacles.")

    # Problem 7
    doc.add_heading("6.7 Problem 7: Gripper Initialization (Critical)", 2)

    p = doc.add_paragraph()
    p.add_run("This was the most critical discovery of the project.").bold = True
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Symptom: ").bold = True
    p.add_run("Robot pushed the cube away during approach (unexpected physics collision).")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Root Cause: ").bold = True
    p.add_run("Gripper was not initialized to default open state before world.reset().")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Solution: ").bold = True
    doc.add_paragraph()

    code = """# CRITICAL: Initialize gripper BEFORE world.reset()
my_franka.gripper.set_default_state(
    my_franka.gripper.joint_opened_positions
)
my_world.reset()  # Now gripper starts in open state"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Result: ").bold = True
    p.add_run("Pick and Place SUCCESS! Distance to target: 0.0071m (7.1mm)")

    # Problem 8
    doc.add_heading("6.8 Problem 8: YCB Collision Avoidance", 2)
    p = doc.add_paragraph()
    p.add_run("Symptom: ").bold = True
    p.add_run("YCB object physics was stable, but gripper stopped 6cm above target (z=0.14m instead of z=0.09m).")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Root Cause: ").bold = True
    p.add_run("RMPFlow conservatively avoided YCB's complex collision mesh, creating large avoidance margin.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Comparison: ").bold = True

    compare_table = doc.add_table(rows=3, cols=2)
    compare_table.style = 'Table Grid'
    comparisons = [
        ("Object Type", "Gripper Z at Phase 4"),
        ("DynamicCuboid", "z = 0.093m (Success)"),
        ("YCB convexDecomposition", "z = 0.193m (6cm too high)"),
    ]
    for i, (obj, z) in enumerate(comparisons):
        row = compare_table.rows[i]
        row.cells[0].text = obj
        row.cells[1].text = z
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Solution: ").bold = True
    p.add_run("Developed the Collision Proxy Pattern (detailed in Section 7).")

    doc.add_page_break()

    # =========================================================================
    # 7. YCB DATASET SUPPORT
    # =========================================================================
    add_heading_with_style(doc, "7. YCB Dataset Support", 1)

    doc.add_paragraph(
        "The YCB (Yale-CMU-Berkeley) Object and Model Set is a standard benchmark for robotic "
        "manipulation research. This project implements support for YCB objects using a novel "
        "Collision Proxy Pattern."
    )

    doc.add_heading("7.1 Collision Proxy Pattern", 2)
    doc.add_paragraph(
        "The key innovation for YCB support is separating visual representation from physics collision:"
    )

    p = doc.add_paragraph()
    p.add_run("Concept: ").bold = True
    p.add_run("Instead of using YCB's complex collision mesh, create an invisible DynamicCuboid as a 'collision proxy' that handles all physics interactions.")

    doc.add_paragraph()
    doc.add_paragraph("Implementation Steps:", style='List Bullet')
    proxy_steps = [
        "Create invisible DynamicCuboid (6cm cube) at object position",
        "Add DynamicCuboid to world.scene (RMPFlow sees this)",
        "Load YCB visual mesh as child of the proxy (no physics)",
        "Proxy handles collision, YCB provides visual appearance",
        "When proxy moves, YCB visual moves with it"
    ]
    for step in proxy_steps:
        doc.add_paragraph(f"  {step}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Code Implementation:").bold = True
    doc.add_paragraph()

    code = """def create_ycb_object(world: World, object_name: str) -> dict:
    # Step 1: Create invisible DynamicCuboid (collision proxy)
    proxy_cube = DynamicCuboid(
        prim_path=f"/World/YCB_Proxy_{object_name}",
        position=OBJECT_POSITION,  # [0.3, 0.3, 0.3]
        scale=np.array([0.06, 0.06, 0.06]),  # 6cm cube
    )
    world.scene.add(proxy_cube)  # RMPFlow sees this

    # Make proxy invisible (physics still active)
    imageable = UsdGeom.Imageable(proxy_prim)
    imageable.MakeInvisible()

    # Step 2: Load YCB visual as child of proxy (no physics)
    ycb_prim_path = f"{proxy_prim_path}/YCB_Visual"
    add_reference_to_stage(usd_path=usd_path, prim_path=ycb_prim_path)
    # NO physics applied - visual only"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_heading("7.2 High-Friction Material", 2)
    doc.add_paragraph(
        "Heavy objects (like Potted Meat Can at 0.35kg) can slip from the gripper during transport. "
        "Solution: Apply high-friction physics material to gripper fingers and object proxy."
    )

    code = """def create_grip_friction_material(stage, material_path, friction=1.0):
    UsdShade.Material.Define(stage, material_path)
    material_prim = stage.GetPrimAtPath(material_path)
    physics_material = UsdPhysics.MaterialAPI.Apply(material_prim)
    physics_material.CreateStaticFrictionAttr(friction)
    physics_material.CreateDynamicFrictionAttr(friction)
    physics_material.CreateRestitutionAttr(0.0)

# Apply to gripper fingers (friction=2.0)
create_grip_friction_material(stage, "/World/Materials/GripperFriction", 2.0)
apply_material_to_prim(stage, f"{franka_path}/panda_leftfinger", mat_path)
apply_material_to_prim(stage, f"{franka_path}/panda_rightfinger", mat_path)"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_heading("7.3 Supported Objects", 2)

    ycb_table = doc.add_table(rows=4, cols=4)
    ycb_table.style = 'Table Grid'
    ycb_objects = [
        ("Object", "YCB ID", "Mass (kg)", "Friction"),
        ("Potted Meat Can", "010_potted_meat_can", "0.35", "2.0"),
        ("Tuna Fish Can", "007_tuna_fish_can", "0.20", "1.0"),
        ("Foam Brick", "061_foam_brick", "0.05", "1.0"),
    ]
    for i, row_data in enumerate(ycb_objects):
        row = ycb_table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_page_break()

    # =========================================================================
    # 8. API REFERENCE
    # =========================================================================
    add_heading_with_style(doc, "8. API Reference", 1)

    doc.add_heading("8.1 Key Imports", 2)

    code = """# Isaac Sim 4.5 New API
from isaacsim.core.api import World
from isaacsim.core.api.scenes import GroundPlane
from isaacsim.core.api.objects import FixedCuboid, DynamicCuboid
from isaacsim.core.prims import SingleXFormPrim

# Robot Manipulators
from isaacsim.robot.manipulators.examples.franka import Franka
from isaacsim.robot.manipulators.examples.franka.controllers import (
    PickPlaceController, RMPFlowController
)

# USD Physics
from pxr import UsdPhysics, PhysxSchema, UsdGeom, UsdShade, Gf

# Utilities
from isaacsim.core.utils.stage import add_reference_to_stage
import numpy as np"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_heading("8.2 VerificationTracker Class", 2)
    doc.add_paragraph(
        "Custom class for tracking and verifying robot actions:"
    )

    code = """class VerificationTracker:
    def __init__(self, world, franka_name, object_prim_path):
        self.world = world
        self.franka_name = franka_name
        self.object_prim_path = object_prim_path
        self.initial_object_position = None
        self.pick_target = None
        self.place_target = None

    def initialize(self):
        \"\"\"Store initial object position.\"\"\"
        self.initial_object_position = self.get_object_position().copy()

    def get_ee_position(self) -> np.ndarray:
        \"\"\"Get current end effector position.\"\"\"
        franka = self.world.scene.get_object(self.franka_name)
        return franka.end_effector.get_world_pose()[0]

    def get_object_position(self) -> np.ndarray:
        \"\"\"Get current object position.\"\"\"
        prim = SingleXFormPrim(self.object_prim_path)
        return prim.get_local_pose()[0]

    def verify_final_result(self) -> bool:
        \"\"\"Check if object reached target container.\"\"\"
        final_pos = self.get_object_position()
        distance = np.linalg.norm(final_pos - self.place_target)
        return distance < 0.05  # 5cm threshold"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_heading("8.3 Configuration Constants", 2)

    code = """# Object configuration
OBJECT_POSITION = np.array([0.3, 0.3, 0.3])
OBJECT_SCALE = np.array([0.0515, 0.0515, 0.0515])

# Place target
CONTAINER_POSITION = np.array([0.3, -0.3, 0.02])
PLACE_POSITION = np.array([0.3, -0.3, 0.27])

# YCB Object Configurations
YCB_OBJECT_CONFIGS = {
    "010_potted_meat_can": {
        "name": "Potted Meat Can",
        "usd_path": "Props/YCB/Axis_Aligned/010_potted_meat_can.usd",
        "mass": 0.35,
        "grip_friction": 2.0,
    },
    # ... more objects
}"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_page_break()

    # =========================================================================
    # 9. TEST RESULTS
    # =========================================================================
    add_heading_with_style(doc, "9. Test Results", 1)

    doc.add_heading("Test History Summary", 2)

    test_table = doc.add_table(rows=8, cols=3)
    test_table.style = 'Table Grid'
    tests = [
        ("Tests", "Focus Area", "Result"),
        ("1-6", "YCB banana, table height adjustments", "Failed"),
        ("7-16", "RMPFlow obstacle investigation", "Failed"),
        ("17-18", "RMPFlow disable_obstacle attempt", "No effect"),
        ("19", "Without table test", "EE XY accurate (0.0147m)"),
        ("20-25", "Pedestal design, VisualCuboid tests", "Partial"),
        ("26", "Gripper initialization added", "SUCCESS!"),
    ]
    for i, (num, focus, result) in enumerate(tests):
        row = test_table.rows[i]
        row.cells[0].text = num
        row.cells[1].text = focus
        row.cells[2].text = result
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        if result == "SUCCESS!":
            set_cell_shading(row.cells[2], "90EE90")

    doc.add_paragraph()
    doc.add_heading("Final Results by Object Type", 2)

    final_table = doc.add_table(rows=5, cols=4)
    final_table.style = 'Table Grid'
    final_results = [
        ("Object", "Phase 4 EE Z", "Distance to Target", "Status"),
        ("DynamicCuboid", "0.093m", "7.1mm", "SUCCESS"),
        ("Potted Meat Can", "0.096m", "11.1mm", "SUCCESS"),
        ("Tuna Fish Can", "0.096m", "11.1mm", "SUCCESS"),
        ("Foam Brick", "0.096m", "11.1mm", "SUCCESS"),
    ]
    for i, row_data in enumerate(final_results):
        row = final_table.rows[i]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, "003366")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        elif row_data[3] == "SUCCESS":
            set_cell_shading(row.cells[3], "90EE90")

    doc.add_page_break()

    # =========================================================================
    # 10. HOW TO RUN
    # =========================================================================
    add_heading_with_style(doc, "10. How to Run", 1)

    doc.add_heading("GUI Mode", 2)

    code = """# Navigate to Isaac Sim directory
cd ~/isaac_sim_4.5

# Run main script with GUI
./python.sh ~/ClaudeCode_PlanMode_PickAndPlace/franka_ycb_pick_place.py

# GUI Mode Controls:
# - Press 'S' to start Pick and Place
# - Press '1' for Potted Meat Can
# - Press '2' for Tuna Fish Can
# - Press '3' for Foam Brick
# - Press 'R' to reset
# - Press 'Q' to quit"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_heading("Headless Mode (Automated Testing)", 2)

    code = """# Run headless test script
cd ~/isaac_sim_4.5
./python.sh ~/ClaudeCode_PlanMode_PickAndPlace/test_headless.py

# Or modify config in franka_ycb_pick_place.py:
CONFIG = {
    "headless": True,  # Set to True for headless mode
    # ...
}"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_heading("Expected Output", 2)

    code = """[FINAL VERIFICATION]
  Initial object position: [0.29995224 0.30002183 0.02574999]
  Final object position:   [ 0.3069876 -0.2988625  0.04975  ]
  Place target position:   [ 0.3  -0.3   0.27]
  Distance moved:          0.5994m
  Distance to target:      0.0071m
[RESULT] SUCCESS - Object is in/near container!

Pick and Place task completed successfully!"""

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(code)
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9)

    doc.add_page_break()

    # =========================================================================
    # 11. FUTURE IMPROVEMENTS
    # =========================================================================
    add_heading_with_style(doc, "11. Future Improvements", 1)

    improvements = [
        ("RMPFlow Configuration Optimization",
         "Tune collision_sphere_buffer and self_collision_avoidance_distance to allow table usage without path planning issues."),
        ("Task Class Refactoring",
         "Migrate to Task-based architecture for cleaner code organization and reusability."),
        ("Multi-Object Scenes",
         "Support picking multiple objects in sequence with obstacle avoidance."),
        ("Grasp Planning",
         "Implement grasp pose optimization for irregularly shaped YCB objects."),
        ("Sensor Integration",
         "Add depth camera feedback for closed-loop manipulation."),
    ]

    for title, desc in improvements:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)
        doc.add_paragraph()

    # =========================================================================
    # 12. REFERENCES
    # =========================================================================
    add_heading_with_style(doc, "12. References", 1)

    doc.add_heading("Project Files", 2)
    refs = [
        "franka_ycb_pick_place.py - Main implementation script",
        "test_headless.py - Automated test script",
        "work_log.txt - Development log with detailed problem analysis",
        "project_presentation_v3.html - Interactive 3D presentation"
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_heading("Isaac Sim Resources", 2)
    resources = [
        "Official Example: ~/isaac_sim_4.5/standalone_examples/api/isaacsim.robot.manipulators/franka_pick_up.py",
        "Isaac Sim Documentation: https://docs.omniverse.nvidia.com/isaacsim/",
        "YCB Dataset: https://www.ycbbenchmarks.com/"
    ]
    for res in resources:
        doc.add_paragraph(res, style='List Bullet')

    # =========================================================================
    # SAVE DOCUMENT
    # =========================================================================
    output_path = "/home/yjchoi/ClaudeCode_PlanMode_PickAndPlace/documents/260119/Isaac_Sim_Pick_and_Place_Technical_Report_v3.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_document()
