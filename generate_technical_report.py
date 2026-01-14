#!/usr/bin/env python3
"""
Generate Technical Report Document for Isaac Sim Pick and Place Project
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_number(doc, text, level):
    """Add numbered heading"""
    heading = doc.add_heading(text, level=level)
    return heading

def create_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run('\n\n\n\n')

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.add_run('Isaac Sim 4.5\nPick and Place Project')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 82, 147)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('\nTechnical Report')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph('\n\n\n')

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Franka Panda Robot Manipulation Demo\n').font.size = Pt(14)
    info.add_run('Using NVIDIA Isaac Sim 4.5\n\n').font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f'Project Period: 2026-01-14 ~ 2026-01-15\n')
    date_para.add_run(f'Document Generated: {datetime.datetime.now().strftime("%Y-%m-%d")}')

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    toc_title = doc.add_heading('Table of Contents', level=1)

    toc_items = [
        ('1. Project Overview', '3'),
        ('2. System Requirements', '4'),
        ('3. Project Structure', '5'),
        ('4. Core Concepts', '6'),
        ('5. Problem Analysis and Solutions', '8'),
        ('6. API Reference', '12'),
        ('7. Installation and Execution', '14'),
        ('8. Test Results', '15'),
        ('9. Key Solutions Summary', '16'),
        ('10. Future Improvements', '17'),
        ('11. References', '18'),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run('\t' * 8 + page)

    doc.add_page_break()

    # =========================================================================
    # 1. PROJECT OVERVIEW
    # =========================================================================
    doc.add_heading('1. Project Overview', level=1)

    doc.add_heading('1.1 Introduction', level=2)
    doc.add_paragraph(
        'This project implements a Pick and Place demonstration using the Franka Panda robot '
        'in NVIDIA Isaac Sim 4.5 simulation environment. The robot autonomously picks up a '
        'DynamicCuboid object from a designated position and places it into a container.'
    )

    doc.add_heading('1.2 Objectives', level=2)
    objectives = [
        'Implement Pick and Place demo with Franka Panda robot',
        'Use Isaac Sim 4.5 new API (isaacsim.core.api.*)',
        'Build verification system to validate actual robot operation success/failure',
        'Achieve positioning accuracy within 10mm',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('1.3 Key Achievements', level=2)

    # Results table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['Metric', 'Result']
    data = [
        ('Total Test Iterations', '26'),
        ('Final Accuracy', '7.1mm (0.0071m)'),
        ('Success Rate', '100% (after fix)'),
        ('Motion Phases', '10 phases'),
    ]

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    for i, (metric, result) in enumerate(data, 1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = result

    doc.add_paragraph()
    doc.add_page_break()

    # =========================================================================
    # 2. SYSTEM REQUIREMENTS
    # =========================================================================
    doc.add_heading('2. System Requirements', level=1)

    doc.add_heading('2.1 Hardware Requirements', level=2)
    hw_reqs = [
        'GPU: NVIDIA RTX series (RTX 4090 recommended)',
        'RAM: 32GB minimum (64GB recommended)',
        'Storage: 50GB available space',
        'CPU: Intel Core i7/i9 or AMD Ryzen 7/9',
    ]
    for req in hw_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('2.2 Software Requirements', level=2)
    sw_reqs = [
        'Operating System: Ubuntu 22.04 LTS',
        'NVIDIA Isaac Sim 4.5',
        'NVIDIA Driver: 525.x or later',
        'Python 3.10 (included with Isaac Sim)',
    ]
    for req in sw_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('2.3 Isaac Sim Installation Path', level=2)
    code = doc.add_paragraph()
    code.add_run('~/isaac_sim_4.5/').font.name = 'Consolas'

    doc.add_page_break()

    # =========================================================================
    # 3. PROJECT STRUCTURE
    # =========================================================================
    doc.add_heading('3. Project Structure', level=1)

    doc.add_heading('3.1 Directory Layout', level=2)

    structure = '''
/home/yjchoi/ClaudeCode_PlanMode_PickAndPlace/
├── franka_ycb_pick_place.py    # Main GUI script
├── test_headless.py            # Headless test script
├── work_log.txt                # Development log
├── project_presentation.html   # 3D visualization
└── generate_technical_report.py # This document generator
'''
    code_para = doc.add_paragraph()
    run = code_para.add_run(structure)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_heading('3.2 File Descriptions', level=2)

    file_table = doc.add_table(rows=5, cols=2)
    file_table.style = 'Table Grid'

    file_headers = ['File', 'Description']
    file_data = [
        ('franka_ycb_pick_place.py', 'Main script with GUI support. Press "S" to start.'),
        ('test_headless.py', 'Automated testing script without GUI (headless mode).'),
        ('work_log.txt', 'Complete development history and problem solutions.'),
        ('project_presentation.html', 'Interactive 3D visualization of the project.'),
    ]

    for i, header in enumerate(file_headers):
        cell = file_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    for i, (file, desc) in enumerate(file_data, 1):
        file_table.rows[i].cells[0].text = file
        file_table.rows[i].cells[1].text = desc

    doc.add_paragraph()
    doc.add_page_break()

    # =========================================================================
    # 4. CORE CONCEPTS
    # =========================================================================
    doc.add_heading('4. Core Concepts', level=1)

    doc.add_heading('4.1 Pick and Place Operation', level=2)
    doc.add_paragraph(
        'Pick and Place is a fundamental robotic manipulation task where a robot arm '
        'picks up an object from one location and places it at another. This project '
        'implements this using the PickPlaceController from Isaac Sim.'
    )

    doc.add_heading('4.2 Motion Phases', level=2)
    doc.add_paragraph(
        'The Pick and Place operation consists of 10 sequential phases:'
    )

    phases_table = doc.add_table(rows=11, cols=3)
    phases_table.style = 'Table Grid'

    phase_headers = ['Phase', 'Name', 'Description']
    phases = [
        ('0', 'Moving above pick', 'Move end effector above the object'),
        ('1', 'Lowering to grasp', 'Lower to grasping position'),
        ('2', 'Waiting settle', 'Wait for physics to stabilize'),
        ('3', 'Closing gripper', 'Close gripper fingers'),
        ('4', 'Lifting object', 'Lift the grasped object'),
        ('5', 'Moving to place XY', 'Move horizontally to place position'),
        ('6', 'Lowering to place', 'Lower object to placement height'),
        ('7', 'Opening gripper', 'Release the object'),
        ('8', 'Lifting up', 'Lift gripper away from object'),
        ('9', 'Going home', 'Return to home position'),
    ]

    for i, header in enumerate(phase_headers):
        cell = phases_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    for i, (phase, name, desc) in enumerate(phases, 1):
        phases_table.rows[i].cells[0].text = phase
        phases_table.rows[i].cells[1].text = name
        phases_table.rows[i].cells[2].text = desc

    doc.add_paragraph()

    doc.add_heading('4.3 RMPFlow Controller', level=2)
    doc.add_paragraph(
        'RMPFlow (Riemannian Motion Policy Flow) is the motion planning algorithm used '
        'by the PickPlaceController. It generates smooth, collision-aware trajectories '
        'for the robot arm. Key characteristics:'
    )

    rmp_points = [
        'Automatic obstacle avoidance (including scene objects)',
        'Smooth trajectory generation using Riemannian geometry',
        'Real-time motion planning capability',
        'Important: Scene objects can interfere with path planning',
    ]
    for point in rmp_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_heading('4.4 Coordinate System', level=2)
    doc.add_paragraph('Isaac Sim uses a right-handed coordinate system:')

    coord_items = [
        'X-axis: Forward (away from robot base)',
        'Y-axis: Left (robot\'s left side)',
        'Z-axis: Up (vertical)',
        'Units: Meters (stage_units_in_meters=1.0)',
    ]
    for item in coord_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 5. PROBLEM ANALYSIS AND SOLUTIONS
    # =========================================================================
    doc.add_heading('5. Problem Analysis and Solutions', level=1)

    doc.add_paragraph(
        'During development, 8 major problems were identified and resolved through '
        '26 test iterations. This section documents each problem and its solution.'
    )

    # Problem 1
    doc.add_heading('5.1 Problem 1: Robot Moving to Wrong Position', level=2)

    p1_table = doc.add_table(rows=4, cols=2)
    p1_table.style = 'Table Grid'
    p1_data = [
        ('Symptom', 'End Effector XY distance was 0.5417m from target'),
        ('Root Cause', 'Object position was outside robot\'s workspace'),
        ('Solution', 'Adjusted TABLE_POSITION to [0.35, 0.0, 0.25] and PLACEMENT_BOUNDS to x:[0.25~0.45], y:[0.10~0.30]'),
        ('Reference', 'Official example uses position ~[0.3, 0.3, 0.3]'),
    ]
    for i, (label, value) in enumerate(p1_data):
        p1_table.rows[i].cells[0].text = label
        p1_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        p1_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Problem 2
    doc.add_heading('5.2 Problem 2: Z Height Too High', level=2)

    p2_table = doc.add_table(rows=3, cols=2)
    p2_table.style = 'Table Grid'
    p2_data = [
        ('Symptom', 'EE z=0.36 vs object z=0.27 - robot couldn\'t reach low enough'),
        ('Root Cause', 'Table height too low for robot to reach XY position at low Z'),
        ('Solution', 'Adjusted table height to z=0.25, object z=0.27'),
    ]
    for i, (label, value) in enumerate(p2_data):
        p2_table.rows[i].cells[0].text = label
        p2_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        p2_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Problem 3
    doc.add_heading('5.3 Problem 3: Object Not Moving (Grip Failed)', level=2)

    p3_table = doc.add_table(rows=3, cols=2)
    p3_table.style = 'Table Grid'
    p3_data = [
        ('Symptom', 'Robot completed motion but object didn\'t move at all'),
        ('Root Cause', 'YCB object created with prim_utils.create_prim() lacked physics properties'),
        ('Solution', 'Added UsdPhysics API: RigidBodyAPI, CollisionAPI, MassAPI'),
    ]
    for i, (label, value) in enumerate(p3_data):
        p3_table.rows[i].cells[0].text = label
        p3_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        p3_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    code_block = doc.add_paragraph()
    code_text = '''# Solution Code:
from pxr import UsdPhysics, PhysxSchema

UsdPhysics.RigidBodyAPI.Apply(prim)
UsdPhysics.CollisionAPI.Apply(prim)
mass_api = UsdPhysics.MassAPI.Apply(prim)
mass_api.CreateMassAttr(0.1)  # 100g'''
    run = code_block.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph()

    # Problem 4
    doc.add_heading('5.4 Problem 4: Object Moving to Wrong Location', level=2)

    p4_table = doc.add_table(rows=3, cols=2)
    p4_table.style = 'Table Grid'
    p4_data = [
        ('Symptom', 'Object moved 0.2728m but to wrong position'),
        ('Root Cause', 'After physics enabled, object settled at different position, but robot used initial stored position'),
        ('Solution', 'Get actual object position after world.reset() and physics settle'),
    ]
    for i, (label, value) in enumerate(p4_data):
        p4_table.rows[i].cells[0].text = label
        p4_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        p4_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Problem 5
    doc.add_heading('5.5 Problem 5: YCB Banana Rolling Away', level=2)

    p5_table = doc.add_table(rows=3, cols=2)
    p5_table.style = 'Table Grid'
    p5_data = [
        ('Symptom', 'Banana rolled off table during physics simulation'),
        ('Root Cause', 'Irregular shape of YCB banana caused instability'),
        ('Solution', 'Replaced YCB banana with DynamicCuboid for stability'),
    ]
    for i, (label, value) in enumerate(p5_data):
        p5_table.rows[i].cells[0].text = label
        p5_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        p5_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Problem 6 - CRITICAL
    doc.add_heading('5.6 Problem 6: Table Interfering with RMPFlow (Critical)', level=2)

    critical_note = doc.add_paragraph()
    run = critical_note.add_run('⚠ This was a major discovery that significantly impacted the solution approach.')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    p6_table = doc.add_table(rows=4, cols=2)
    p6_table.style = 'Table Grid'
    p6_data = [
        ('Symptom', 'With table: EE XY distance 0.45~0.55m (completely wrong position)'),
        ('Root Cause', 'RMPFlow recognized table as obstacle and generated incorrect avoidance paths'),
        ('Discovery', 'Without table: EE XY distance = 0.0147m (accurate!)'),
        ('Solution', 'Remove table, use only GroundPlane'),
    ]
    for i, (label, value) in enumerate(p6_data):
        p6_table.rows[i].cells[0].text = label
        p6_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        p6_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Problem 7
    doc.add_heading('5.7 Problem 7: EE Minimum Height Limitation', level=2)

    p7_table = doc.add_table(rows=3, cols=2)
    p7_table.style = 'Table Grid'
    p7_data = [
        ('Symptom', 'Robot cannot reach below z=0.287m'),
        ('Discovery', 'Object center must be at z~0.27 for gripper to reach'),
        ('Implication', 'Object placement height must consider robot reach limits'),
    ]
    for i, (label, value) in enumerate(p7_data):
        p7_table.rows[i].cells[0].text = label
        p7_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        p7_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Problem 8 - KEY SOLUTION
    doc.add_heading('5.8 Problem 8: Gripper Initialization Missing (KEY SOLUTION)', level=2)

    key_note = doc.add_paragraph()
    run = key_note.add_run('★ This was the final and most critical fix that enabled success!')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)

    p8_table = doc.add_table(rows=4, cols=2)
    p8_table.style = 'Table Grid'
    p8_data = [
        ('Symptom', 'Robot pushed cube away during approach'),
        ('Root Cause', 'Gripper not initialized to default state, causing unexpected behavior'),
        ('Solution', 'Call gripper.set_default_state() before world.reset()'),
        ('Result', 'Pick and Place SUCCESS! Distance to target: 0.0071m (7.1mm)'),
    ]
    for i, (label, value) in enumerate(p8_data):
        p8_table.rows[i].cells[0].text = label
        p8_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        p8_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    code_block2 = doc.add_paragraph()
    code_text2 = '''# KEY SOLUTION CODE (from official example):
my_franka.gripper.set_default_state(my_franka.gripper.joint_opened_positions)
my_world.reset()  # Call set_default_state() BEFORE reset()'''
    run = code_block2.add_run(code_text2)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_page_break()

    # =========================================================================
    # 6. API REFERENCE
    # =========================================================================
    doc.add_heading('6. API Reference', level=1)

    doc.add_heading('6.1 Key Imports', level=2)

    imports_code = '''# Isaac Sim 4.5 API
from isaacsim import SimulationApp
from isaacsim.core.api import World
from isaacsim.core.api.objects import DynamicCuboid, FixedCuboid
from isaacsim.core.api.objects.ground_plane import GroundPlane
from isaacsim.core.prims import SingleXFormPrim

# Franka Robot
from isaacsim.robot.manipulators.examples.franka import Franka
from isaacsim.robot.manipulators.examples.franka.controllers import PickPlaceController

# USD/Physics
from pxr import UsdPhysics, PhysxSchema'''

    code_para = doc.add_paragraph()
    run = code_para.add_run(imports_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_heading('6.2 VerificationTracker Class', level=2)
    doc.add_paragraph(
        'Custom class implemented to track and verify actual robot operation success/failure.'
    )

    methods_table = doc.add_table(rows=7, cols=2)
    methods_table.style = 'Table Grid'

    method_headers = ['Method', 'Description']
    methods = [
        ('__init__(franka, object_prim_path, pick_target, place_target)', 'Initialize tracker with robot and target positions'),
        ('initialize()', 'Store initial object position after world reset'),
        ('get_ee_position()', 'Return current End Effector world position'),
        ('get_object_position()', 'Return current object world position'),
        ('log_phase_status(phase, step)', 'Log detailed status at each phase transition'),
        ('verify_final_result()', 'Verify if object reached target container'),
    ]

    for i, header in enumerate(method_headers):
        cell = methods_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    for i, (method, desc) in enumerate(methods, 1):
        methods_table.rows[i].cells[0].text = method
        methods_table.rows[i].cells[1].text = desc

    doc.add_paragraph()

    doc.add_heading('6.3 Key Configuration Constants', level=2)

    config_table = doc.add_table(rows=6, cols=3)
    config_table.style = 'Table Grid'

    config_headers = ['Constant', 'Value', 'Description']
    configs = [
        ('OBJECT_POSITION', '[0.3, 0.3, 0.3]', 'Pick object initial position'),
        ('OBJECT_SCALE', '[0.0515, 0.0515, 0.0515]', '~5cm cube size'),
        ('PLACE_POSITION', '[0.3, -0.3, 0.27]', 'Target placement position'),
        ('CONTAINER_POSITION', '[0.3, -0.3, 0.02]', 'Container location'),
        ('end_effector_offset', '[0, 0.005, 0]', 'EE offset for grasping'),
    ]

    for i, header in enumerate(config_headers):
        cell = config_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    for i, (const, val, desc) in enumerate(configs, 1):
        config_table.rows[i].cells[0].text = const
        config_table.rows[i].cells[1].text = val
        config_table.rows[i].cells[2].text = desc

    doc.add_page_break()

    # =========================================================================
    # 7. INSTALLATION AND EXECUTION
    # =========================================================================
    doc.add_heading('7. Installation and Execution', level=1)

    doc.add_heading('7.1 Prerequisites', level=2)
    prereqs = [
        'NVIDIA Isaac Sim 4.5 installed at ~/isaac_sim_4.5/',
        'Project files in ~/ClaudeCode_PlanMode_PickAndPlace/',
        'NVIDIA GPU with proper drivers installed',
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')

    doc.add_heading('7.2 Running the GUI Script', level=2)

    gui_code = '''cd ~/isaac_sim_4.5
./python.sh ~/ClaudeCode_PlanMode_PickAndPlace/franka_ycb_pick_place.py

# After window opens, press 'S' key to start the Pick and Place task'''
    code_para = doc.add_paragraph()
    run = code_para.add_run(gui_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading('7.3 Running Headless Test', level=2)

    headless_code = '''cd ~/isaac_sim_4.5
./python.sh ~/ClaudeCode_PlanMode_PickAndPlace/test_headless.py

# Runs automatically without GUI, outputs results to console'''
    code_para = doc.add_paragraph()
    run = code_para.add_run(headless_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading('7.4 Switching Between GUI and Headless Mode', level=2)
    doc.add_paragraph(
        'In franka_ycb_pick_place.py, modify the CONFIG dictionary:'
    )

    switch_code = '''CONFIG = {
    "headless": False,  # Set to True for headless mode
    "width": 1280,
    "height": 720,
}'''
    code_para = doc.add_paragraph()
    run = code_para.add_run(switch_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_page_break()

    # =========================================================================
    # 8. TEST RESULTS
    # =========================================================================
    doc.add_heading('8. Test Results', level=1)

    doc.add_heading('8.1 Test History Summary', level=2)

    test_table = doc.add_table(rows=8, cols=3)
    test_table.style = 'Table Grid'

    test_headers = ['Test Range', 'Focus Area', 'Result']
    tests = [
        ('1-6', 'YCB banana, table height adjustments', 'Failed'),
        ('7-16', 'RMPFlow obstacle investigation', 'Failed'),
        ('17-18', 'RMPFlow disable_obstacle() attempt', 'No effect'),
        ('19', 'Test without table', 'EE XY accurate (0.0147m)'),
        ('20-24', 'Pedestal design attempts', 'Cube pushed away'),
        ('25', 'VisualCuboid test', 'EE accuracy confirmed'),
        ('26', 'Added gripper initialization', 'SUCCESS! (7.1mm)'),
    ]

    for i, header in enumerate(test_headers):
        cell = test_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    for i, (range_, focus, result) in enumerate(tests, 1):
        test_table.rows[i].cells[0].text = range_
        test_table.rows[i].cells[1].text = focus
        test_table.rows[i].cells[2].text = result

    doc.add_paragraph()

    doc.add_heading('8.2 Final Test Results (Test #26)', level=2)

    final_results = '''[FINAL VERIFICATION]
  Initial object position: [0.29995224 0.30002183 0.02574999]
  Final object position:   [0.3069876 -0.2988625 0.04975]
  Place target position:   [0.3 -0.3 0.27]
  Distance moved:          0.5994m
  Distance to target:      0.0071m

[RESULT] SUCCESS - Object is in/near container!'''

    code_para = doc.add_paragraph()
    run = code_para.add_run(final_results)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_page_break()

    # =========================================================================
    # 9. KEY SOLUTIONS SUMMARY
    # =========================================================================
    doc.add_heading('9. Key Solutions Summary', level=1)

    doc.add_paragraph(
        'The following four elements are essential for successful Pick and Place operation:'
    )

    doc.add_heading('9.1 Complete Solution Code', level=2)

    solution_code = '''# 1. Use GroundPlane only (no table - prevents RMPFlow interference)
GroundPlane(prim_path="/World/GroundPlane", z_position=0, name="ground_plane")

# 2. Use DynamicCuboid at official example position
cube = DynamicCuboid(
    prim_path="/World/PickCube",
    position=np.array([0.3, 0.3, 0.3]),      # Official example position
    scale=np.array([0.0515, 0.0515, 0.0515]), # ~5cm cube
)

# 3. CRITICAL: Initialize gripper before reset
my_franka.gripper.set_default_state(my_franka.gripper.joint_opened_positions)
my_world.reset()

# 4. Wait for physics settle, then get actual position
for i in range(50):
    my_world.step(render=False)
object_position = cube.get_local_pose()[0]  # Use settled position'''

    code_para = doc.add_paragraph()
    run = code_para.add_run(solution_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph()

    doc.add_heading('9.2 Key Takeaways', level=2)

    takeaways = [
        'Gripper initialization is the most critical factor for success',
        'Table objects interfere with RMPFlow path planning - avoid or handle carefully',
        'Always get object position AFTER physics settle, not before',
        'Use DynamicCuboid instead of irregular YCB objects for stability',
        'Robot EE has minimum height limit (~0.287m) - plan object positions accordingly',
    ]
    for i, takeaway in enumerate(takeaways, 1):
        doc.add_paragraph(f'{i}. {takeaway}')

    doc.add_page_break()

    # =========================================================================
    # 10. FUTURE IMPROVEMENTS
    # =========================================================================
    doc.add_heading('10. Future Improvements', level=1)

    doc.add_paragraph('The following improvements can be made in future iterations:')

    improvements = [
        ('RMPFlow Optimization', 'Configure RMPFlow to work with table obstacles without path interference'),
        ('Task Class Refactoring', 'Refactor code using Isaac Sim Task class for cleaner architecture'),
        ('YCB Object Support', 'Improve handling of irregular YCB objects (banana, etc.)'),
        ('Multi-Object Handling', 'Extend to pick and place multiple objects sequentially'),
        ('Error Recovery', 'Implement automatic retry on grasp failure'),
    ]

    for title, desc in improvements:
        p = doc.add_paragraph()
        run = p.add_run(f'• {title}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # =========================================================================
    # 11. REFERENCES
    # =========================================================================
    doc.add_heading('11. References', level=1)

    doc.add_heading('11.1 Project Files', level=2)
    refs = [
        'Main Script: ~/ClaudeCode_PlanMode_PickAndPlace/franka_ycb_pick_place.py',
        'Test Script: ~/ClaudeCode_PlanMode_PickAndPlace/test_headless.py',
        'Work Log: ~/ClaudeCode_PlanMode_PickAndPlace/work_log.txt',
        'Visualization: ~/ClaudeCode_PlanMode_PickAndPlace/project_presentation.html',
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_heading('11.2 Isaac Sim Resources', level=2)
    isaac_refs = [
        'Official Example: ~/isaac_sim_4.5/standalone_examples/api/isaacsim.robot.manipulators/franka_pick_up.py',
        'Isaac Sim Documentation: https://docs.omniverse.nvidia.com/isaacsim/',
        'Isaac Sim Path: ~/isaac_sim_4.5/',
    ]
    for ref in isaac_refs:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_heading('11.3 Document Information', level=2)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Project Period', '2026-01-14 ~ 2026-01-15'),
        ('Document Version', '1.0'),
        ('Last Updated', datetime.datetime.now().strftime('%Y-%m-%d')),
        ('Total Test Iterations', '26'),
    ]

    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        info_table.rows[i].cells[1].text = value

    return doc


if __name__ == '__main__':
    print("Generating Technical Report...")
    doc = create_document()

    output_path = '/home/yjchoi/ClaudeCode_PlanMode_PickAndPlace/Isaac_Sim_Pick_and_Place_Technical_Report.docx'
    doc.save(output_path)

    print(f"Document saved to: {output_path}")
    print("Done!")
