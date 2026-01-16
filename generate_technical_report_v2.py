#!/usr/bin/env python3
"""
Generate Technical Report for Isaac Sim Pick and Place Project
Creates a comprehensive .docx document covering the entire project journey.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)

def create_document():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================

    # Add some spacing at top
    for _ in range(3):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Isaac Sim 4.5\nPick and Place Project")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(79, 70, 229)  # Indigo

    doc.add_paragraph()

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Technical Report")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Franka Panda 로봇을 사용한 자율 물체 조작 시뮬레이션\n"
                       "YCB 데이터셋 객체 Pick and Place 구현")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(120, 120, 120)

    for _ in range(5):
        doc.add_paragraph()

    # Key metrics
    metrics = doc.add_paragraph()
    metrics.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = metrics.add_run("26회 테스트  |  8개 문제 해결  |  7.1mm 정확도")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)  # Green

    for _ in range(8):
        doc.add_paragraph()

    # Date and info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"작성일: 2026년 1월 16일\n"
                       f"프로젝트 기간: 2026.01.14 - 2026.01.16")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)

    # Page break
    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================

    doc.add_heading('목차 (Table of Contents)', level=1)

    toc_items = [
        ("1. 프로젝트 개요 (Project Overview)", 3),
        ("2. 시스템 요구사항 (System Requirements)", 4),
        ("3. 프로젝트 구조 (Project Structure)", 5),
        ("4. 핵심 개념 (Core Concepts)", 6),
        ("   4.1 Pick and Place 동작 원리", 6),
        ("   4.2 10단계 Phase 시퀀스", 6),
        ("   4.3 RMPFlow Motion Planner", 7),
        ("5. 문제 분석 및 해결 (Problem Analysis & Solutions)", 8),
        ("   5.1 Problem #1-4: 초기 설정 문제", 8),
        ("   5.2 Problem #5-7: RMPFlow 충돌 문제", 9),
        ("   5.3 Problem #8: 그리퍼 초기화 (핵심 해결책)", 10),
        ("6. YCB 데이터셋 지원 (YCB Dataset Support)", 11),
        ("   6.1 YCB 객체 문제점", 11),
        ("   6.2 Collision Proxy 패턴 (최종 해결책)", 12),
        ("7. API 레퍼런스 (API Reference)", 13),
        ("8. 테스트 결과 (Test Results)", 15),
        ("9. 실행 방법 (How to Run)", 16),
        ("10. 향후 개선 사항 (Future Improvements)", 17),
        ("11. 참고 자료 (References)", 18),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run("\t" * 8 + str(page))

    doc.add_page_break()

    # =========================================================================
    # 1. PROJECT OVERVIEW
    # =========================================================================

    doc.add_heading('1. 프로젝트 개요 (Project Overview)', level=1)

    doc.add_heading('1.1 프로젝트 목표', level=2)
    doc.add_paragraph(
        "이 프로젝트의 목표는 NVIDIA Isaac Sim 4.5 시뮬레이터에서 Franka Panda 로봇 팔을 사용하여 "
        "물체를 집어서(Pick) 지정된 위치에 놓는(Place) 자율 조작 데모를 구현하는 것입니다."
    )

    objectives = [
        "Franka Panda 로봇의 Pick and Place 동작 구현",
        "Isaac Sim 4.5 새로운 API (isaacsim.core.api.*) 활용",
        "YCB 데이터셋 객체 지원 (다양한 형태의 물체)",
        "로봇 동작의 실제 성공/실패를 검증하는 시스템 구현",
        "7.1mm 이하의 배치 정확도 달성",
    ]

    for obj in objectives:
        p = doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading('1.2 프로젝트 성과', level=2)

    # Results table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    results = [
        ("총 테스트 횟수", "26회"),
        ("해결한 문제 수", "8개"),
        ("DynamicCuboid 정확도", "7.1mm"),
        ("YCB 객체 정확도", "11.1mm"),
        ("프로젝트 기간", "3일 (2026.01.14-16)"),
    ]

    for i, (key, value) in enumerate(results):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        set_cell_shading(row.cells[0], 'E5E7EB')

    doc.add_paragraph()

    doc.add_heading('1.3 기술 스택', level=2)

    tech_stack = [
        ("시뮬레이터", "NVIDIA Isaac Sim 4.5"),
        ("로봇", "Franka Panda (7-DOF 매니퓰레이터)"),
        ("Motion Planner", "RMPFlow (Riemannian Motion Policy)"),
        ("프로그래밍 언어", "Python 3.10"),
        ("3D 프레임워크", "USD (Universal Scene Description)"),
        ("Physics Engine", "NVIDIA PhysX"),
    ]

    for name, desc in tech_stack:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # =========================================================================
    # 2. SYSTEM REQUIREMENTS
    # =========================================================================

    doc.add_heading('2. 시스템 요구사항 (System Requirements)', level=1)

    doc.add_heading('2.1 하드웨어 요구사항', level=2)

    hw_reqs = [
        ("GPU", "NVIDIA RTX 3080 이상 (RTX 4090 권장)"),
        ("VRAM", "최소 8GB, 권장 16GB 이상"),
        ("RAM", "최소 32GB, 권장 64GB"),
        ("저장 공간", "50GB 이상 (Isaac Sim + Assets)"),
        ("CPU", "8코어 이상 (Intel i9-12900K 권장)"),
    ]

    table = doc.add_table(rows=len(hw_reqs), cols=2)
    table.style = 'Table Grid'

    for i, (comp, spec) in enumerate(hw_reqs):
        table.rows[i].cells[0].text = comp
        table.rows[i].cells[1].text = spec
        set_cell_shading(table.rows[i].cells[0], 'E5E7EB')

    doc.add_paragraph()

    doc.add_heading('2.2 소프트웨어 요구사항', level=2)

    sw_reqs = [
        "Ubuntu 22.04 LTS",
        "NVIDIA Driver 535+ (CUDA 12.x 지원)",
        "NVIDIA Isaac Sim 4.5",
        "Python 3.10 (Isaac Sim 내장)",
    ]

    for req in sw_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('2.3 Python 패키지 (Isaac Sim 내장)', level=2)

    packages = [
        ("isaacsim.core.api", "핵심 API (World, Scene, Objects)"),
        ("isaacsim.robot.manipulators", "로봇 매니퓰레이터"),
        ("isaacsim.core.prims", "USD Prim 래퍼"),
        ("pxr (USD)", "Universal Scene Description"),
        ("numpy", "수치 연산"),
    ]

    for pkg, desc in packages:
        p = doc.add_paragraph()
        run = p.add_run(f"{pkg}")
        run.bold = True
        run.font.name = 'Consolas'
        p.add_run(f" - {desc}")

    doc.add_page_break()

    # =========================================================================
    # 3. PROJECT STRUCTURE
    # =========================================================================

    doc.add_heading('3. 프로젝트 구조 (Project Structure)', level=1)

    doc.add_heading('3.1 디렉토리 구조', level=2)

    structure = """
ClaudeCode_PlanMode_PickAndPlace/
├── franka_ycb_pick_place.py      # 메인 스크립트 (GUI/Headless)
├── test_headless.py              # Headless 테스트 스크립트
├── work_log.txt                  # 작업 로그
├── README.md                     # 프로젝트 소개
├── project_presentation_v2.html  # 3D 애니메이션 프레젠테이션
├── generate_technical_report_v2.py  # 이 문서 생성 스크립트
├── .gitignore                    # Git 제외 패턴
└── documents/                    # 문서 및 로그 폴더
    ├── 260114/                   # 1월 14일 작업 파일
    ├── 260115_1/                 # 1월 15일 세션 1
    └── 260115_2/                 # 1월 15일 세션 2
    """

    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('3.2 주요 파일 설명', level=2)

    files = [
        ("franka_ycb_pick_place.py",
         "메인 실행 스크립트. GUI 모드와 Headless 모드를 모두 지원합니다. "
         "Franka 로봇, 컨테이너, 객체(DynamicCuboid 또는 YCB) 생성 및 "
         "PickPlaceController를 사용한 Pick and Place 동작을 수행합니다."),

        ("test_headless.py",
         "자동화된 테스트용 Headless 스크립트. GUI 없이 빠르게 테스트를 수행하고 "
         "결과를 검증합니다."),

        ("work_log.txt",
         "프로젝트 진행 과정의 상세한 작업 로그. 발견한 문제, 해결 과정, "
         "테스트 결과 등이 기록되어 있습니다."),
    ]

    for filename, desc in files:
        p = doc.add_paragraph()
        run = p.add_run(f"{filename}")
        run.bold = True
        run.font.name = 'Consolas'
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =========================================================================
    # 4. CORE CONCEPTS
    # =========================================================================

    doc.add_heading('4. 핵심 개념 (Core Concepts)', level=1)

    doc.add_heading('4.1 Pick and Place 동작 원리', level=2)

    doc.add_paragraph(
        "Pick and Place는 로봇 공학에서 가장 기본적인 조작 작업입니다. "
        "로봇 팔이 특정 위치의 물체를 집어서(Pick) 다른 위치에 놓는(Place) 동작을 수행합니다."
    )

    doc.add_paragraph(
        "이 프로젝트에서는 Isaac Sim의 PickPlaceController를 사용하여 "
        "10단계의 정밀한 동작 시퀀스를 통해 Pick and Place를 구현합니다."
    )

    # Coordinate system
    doc.add_heading('좌표계', level=3)

    coords = [
        "X축: 로봇 전방 방향 (양수: 전방)",
        "Y축: 로봇 좌측 방향 (양수: 좌측)",
        "Z축: 수직 방향 (양수: 위쪽)",
        "원점: 로봇 베이스 중심",
    ]

    for c in coords:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_heading('4.2 10단계 Phase 시퀀스', level=2)

    doc.add_paragraph(
        "PickPlaceController는 Pick and Place 동작을 10개의 Phase로 나누어 순차적으로 실행합니다."
    )

    phases = [
        ("Phase 0", "Moving above pick", "Pick 위치 위로 이동"),
        ("Phase 1", "Lowering to grasp", "Grasp 위치로 하강"),
        ("Phase 2", "Waiting settle", "안정화 대기"),
        ("Phase 3", "Closing gripper", "그리퍼 닫기 (잡기)"),
        ("Phase 4", "Lifting object", "객체 들어올리기"),
        ("Phase 5", "Moving to place XY", "Place 위치 XY로 이동"),
        ("Phase 6", "Lowering to place", "Place 위치로 하강"),
        ("Phase 7", "Opening gripper", "그리퍼 열기 (놓기)"),
        ("Phase 8", "Lifting up", "위로 들어올리기"),
        ("Phase 9", "Going home", "홈 위치로 복귀"),
    ]

    table = doc.add_table(rows=len(phases)+1, cols=3)
    table.style = 'Table Grid'

    # Header
    header = table.rows[0]
    header.cells[0].text = "Phase"
    header.cells[1].text = "Name"
    header.cells[2].text = "Description"
    for cell in header.cells:
        set_cell_shading(cell, '4F46E5')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True

    for i, (phase, name, desc) in enumerate(phases):
        row = table.rows[i+1]
        row.cells[0].text = phase
        row.cells[1].text = name
        row.cells[2].text = desc

    doc.add_paragraph()

    doc.add_heading('4.3 RMPFlow Motion Planner', level=2)

    doc.add_paragraph(
        "RMPFlow (Riemannian Motion Policy Flow)는 NVIDIA에서 개발한 반응형 모션 플래너입니다. "
        "실시간으로 장애물을 회피하면서 목표 위치로 이동하는 경로를 생성합니다."
    )

    rmp_features = [
        "실시간 장애물 회피 (Collision Avoidance)",
        "관절 제한 준수 (Joint Limits)",
        "부드러운 경로 생성 (Smooth Trajectories)",
        "자기 충돌 방지 (Self-Collision Avoidance)",
    ]

    for feat in rmp_features:
        doc.add_paragraph(feat, style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run("주의: ")
    run.bold = True
    run.font.color.rgb = RGBColor(239, 68, 68)
    p.add_run("RMPFlow는 scene의 collision geometry를 장애물로 인식합니다. "
              "복잡한 mesh collision은 과도한 회피 동작을 유발할 수 있습니다.")

    doc.add_page_break()

    # =========================================================================
    # 5. PROBLEM ANALYSIS & SOLUTIONS
    # =========================================================================

    doc.add_heading('5. 문제 분석 및 해결 (Problem Analysis & Solutions)', level=1)

    doc.add_paragraph(
        "프로젝트 진행 중 총 8가지 주요 문제를 발견하고 해결했습니다. "
        "각 문제의 증상, 원인, 해결 방법을 상세히 기록합니다."
    )

    doc.add_heading('5.1 Problem #1-4: 초기 설정 문제', level=2)

    # Problem 1
    doc.add_heading('Problem #1: 로봇이 완전히 잘못된 위치로 이동', level=3)

    p = doc.add_paragraph()
    run = p.add_run("증상: ")
    run.bold = True
    p.add_run("End Effector XY 거리가 0.5417m - 객체와 전혀 다른 위치로 이동")

    p = doc.add_paragraph()
    run = p.add_run("원인: ")
    run.bold = True
    p.add_run("객체 위치가 Franka 로봇의 작업 범위(Workspace) 밖에 있었음")

    p = doc.add_paragraph()
    run = p.add_run("해결: ")
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    p.add_run("객체 위치를 [0.3, 0.3, 0.3]으로 조정 (공식 예제 참고)")

    add_horizontal_line(doc)

    # Problem 2
    doc.add_heading('Problem #2: Z 높이가 너무 높음', level=3)

    p = doc.add_paragraph()
    run = p.add_run("증상: ")
    run.bold = True
    p.add_run("EE z=0.36 vs 객체 z=0.27 - 로봇이 충분히 내려가지 못함")

    p = doc.add_paragraph()
    run = p.add_run("원인: ")
    run.bold = True
    p.add_run("테이블 높이가 너무 낮아 로봇이 해당 XY에서 낮은 Z에 도달 불가")

    p = doc.add_paragraph()
    run = p.add_run("해결: ")
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    p.add_run("테이블 높이를 z=0.25로 조정, 객체를 z=0.27에 배치")

    add_horizontal_line(doc)

    # Problem 3
    doc.add_heading('Problem #3: 객체가 움직이지 않음', level=3)

    p = doc.add_paragraph()
    run = p.add_run("증상: ")
    run.bold = True
    p.add_run("로봇 동작은 완료되지만 객체가 전혀 이동하지 않음")

    p = doc.add_paragraph()
    run = p.add_run("원인: ")
    run.bold = True
    p.add_run("YCB 객체에 Physics 속성(RigidBody, Collision)이 없었음")

    p = doc.add_paragraph()
    run = p.add_run("해결: ")
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    p.add_run("UsdPhysics API 추가")

    code = """
from pxr import UsdPhysics

UsdPhysics.RigidBodyAPI.Apply(prim)
UsdPhysics.CollisionAPI.Apply(prim)
mass_api = UsdPhysics.MassAPI.Apply(prim)
mass_api.CreateMassAttr(0.1)  # 100g
"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    add_horizontal_line(doc)

    # Problem 4
    doc.add_heading('Problem #4: 객체가 잘못된 위치로 이동', level=3)

    p = doc.add_paragraph()
    run = p.add_run("증상: ")
    run.bold = True
    p.add_run("객체는 움직이지만 목표가 아닌 곳으로 이동 (distance: 0.27m)")

    p = doc.add_paragraph()
    run = p.add_run("원인: ")
    run.bold = True
    p.add_run("Physics settle 후 객체 위치가 변경되었는데, 로봇은 초기 저장된 위치를 타겟팅")

    p = doc.add_paragraph()
    run = p.add_run("해결: ")
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    p.add_run("world.reset() 후 객체의 실제 위치를 다시 획득")

    code = """
# Physics settle 후 객체의 실제 위치 획득
object_position = verifier.get_object_position().copy()
"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_page_break()

    doc.add_heading('5.2 Problem #5-7: RMPFlow 충돌 문제', level=2)

    # Problem 5
    doc.add_heading('Problem #5: YCB 객체가 굴러감', level=3)

    p = doc.add_paragraph()
    run = p.add_run("증상: ")
    run.bold = True
    p.add_run("바나나가 physics settle 중 테이블에서 떨어짐")

    p = doc.add_paragraph()
    run = p.add_run("원인: ")
    run.bold = True
    p.add_run("YCB 바나나의 불규칙한 모양으로 인해 테이블 위에서 불안정")

    p = doc.add_paragraph()
    run = p.add_run("해결: ")
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    p.add_run("DynamicCuboid 사용 (안정적인 정육면체 형태)")

    add_horizontal_line(doc)

    # Problem 6
    doc.add_heading('Problem #6: 테이블이 RMPFlow 경로를 방해', level=3)

    p = doc.add_paragraph()
    run = p.add_run("증상: ")
    run.bold = True
    p.add_run("테이블이 있으면 EE XY 거리 0.45~0.55m (완전히 잘못된 경로)")

    p = doc.add_paragraph()
    run = p.add_run("원인: ")
    run.bold = True
    p.add_run("RMPFlow가 테이블을 장애물로 인식하여 회피 경로 생성")

    p = doc.add_paragraph()
    run = p.add_run("발견: ")
    run.bold = True

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "조건"
    table.rows[0].cells[1].text = "EE XY 거리"
    table.rows[1].cells[0].text = "테이블 없음"
    table.rows[1].cells[1].text = "0.0147m (정확!)"
    table.rows[2].cells[0].text = "테이블 있음"
    table.rows[2].cells[1].text = "0.45~0.55m (실패)"
    set_cell_shading(table.rows[0].cells[0], 'E5E7EB')
    set_cell_shading(table.rows[0].cells[1], 'E5E7EB')

    p = doc.add_paragraph()
    run = p.add_run("해결: ")
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    p.add_run("테이블 제거, GroundPlane만 사용")

    add_horizontal_line(doc)

    # Problem 7
    doc.add_heading('Problem #7: EE 최저 높이 제한', level=3)

    p = doc.add_paragraph()
    run = p.add_run("증상: ")
    run.bold = True
    p.add_run("로봇이 z=0.287m 이하로 내려가지 못함")

    p = doc.add_paragraph()
    run = p.add_run("원인: ")
    run.bold = True
    p.add_run("Franka 로봇의 기구학적 제한 및 충돌 회피")

    p = doc.add_paragraph()
    run = p.add_run("해결: ")
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    p.add_run("객체를 z~0.27에 배치 (EE 도달 가능 높이)")

    doc.add_page_break()

    doc.add_heading('5.3 Problem #8: 그리퍼 초기화 (핵심 해결책)', level=2)

    p = doc.add_paragraph()
    run = p.add_run("이것이 프로젝트의 가장 중요한 해결책입니다!")
    run.bold = True
    run.font.color.rgb = RGBColor(239, 68, 68)
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("증상: ")
    run.bold = True
    p.add_run("로봇이 객체에 접근하는 중에 객체를 밀어냄 (physics collision)")

    p = doc.add_paragraph()
    run = p.add_run("원인: ")
    run.bold = True
    p.add_run("그리퍼가 기본 상태(열린 상태)로 초기화되지 않아 예상치 못한 동작 발생")

    p = doc.add_paragraph()
    run = p.add_run("해결: ")
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)
    p.add_run("world.reset() 전에 그리퍼 기본 상태 설정")

    code = """
# 핵심 코드 - 그리퍼 초기화
my_franka.gripper.set_default_state(my_franka.gripper.joint_opened_positions)
my_world.reset()  # reset() 전에 set_default_state() 호출!
"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("결과: ")
    run.bold = True
    p.add_run("Pick and Place 성공! 타겟까지 거리 ")
    run = p.add_run("0.0071m (7.1mm)")
    run.bold = True
    run.font.color.rgb = RGBColor(16, 185, 129)

    doc.add_page_break()

    # =========================================================================
    # 6. YCB DATASET SUPPORT
    # =========================================================================

    doc.add_heading('6. YCB 데이터셋 지원 (YCB Dataset Support)', level=1)

    doc.add_paragraph(
        "YCB (Yale-CMU-Berkeley) 데이터셋은 로봇 조작 연구를 위한 표준 객체 세트입니다. "
        "다양한 형태, 크기, 재질의 일상 물체들을 포함하고 있습니다."
    )

    doc.add_heading('6.1 YCB 객체 문제점', level=2)

    doc.add_paragraph(
        "DynamicCuboid로 성공한 후, YCB 데이터셋 객체로 확장을 시도했습니다. "
        "그러나 새로운 문제가 발생했습니다."
    )

    p = doc.add_paragraph()
    run = p.add_run("문제: ")
    run.bold = True
    p.add_run("그리퍼가 YCB 객체 위 6cm에서 정지 (도달 실패)")

    # Comparison table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    headers = ["객체", "Phase 4 EE Z", "결과"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], '4F46E5')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    table.rows[1].cells[0].text = "DynamicCuboid"
    table.rows[1].cells[1].text = "z = 0.093m"
    table.rows[1].cells[2].text = "SUCCESS"

    table.rows[2].cells[0].text = "YCB 객체"
    table.rows[2].cells[1].text = "z = 0.14m"
    table.rows[2].cells[2].text = "FAILED"

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("원인 분석: ")
    run.bold = True
    p.add_run("RMPFlow가 YCB의 복잡한 collision mesh를 장애물로 인식하여 "
              "과도한 회피 margin(~6cm)을 적용")

    doc.add_heading('6.2 Collision Proxy 패턴 (최종 해결책)', level=2)

    p = doc.add_paragraph()
    run = p.add_run("핵심 아이디어: ")
    run.bold = True
    p.add_run("YCB 객체의 복잡한 collision mesh 대신, 보이지 않는 DynamicCuboid를 "
              "'collision proxy'로 사용합니다.")

    doc.add_paragraph()
    doc.add_heading('구현 방법', level=3)

    steps = [
        "보이지 않는 DynamicCuboid (6cm cube) 생성 - physics/collision 담당",
        "YCB 시각 메시를 proxy의 자식으로 로드 - 시각적 표현만 (physics 없음)",
        "Proxy cube가 DynamicCuboid와 동일하게 동작하므로 gripper 접근 가능",
        "YCB visual은 proxy와 함께 이동",
    ]

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. ")
        run.bold = True
        p.add_run(step)

    doc.add_paragraph()
    doc.add_heading('핵심 코드', level=3)

    code = """
def create_ycb_object(world: World, object_name: str) -> dict:
    # Step 1: 보이지 않는 DynamicCuboid 생성 (collision proxy)
    proxy_cube = DynamicCuboid(
        prim_path=f"/World/YCB_Proxy_{object_name}",
        position=OBJECT_POSITION,  # [0.3, 0.3, 0.3]
        scale=np.array([0.06, 0.06, 0.06]),  # 6cm cube
        color=np.array([1, 1, 1]),
    )
    world.scene.add(proxy_cube)

    # Proxy cube를 보이지 않게 설정 (physics는 유지)
    imageable = UsdGeom.Imageable(proxy_prim)
    imageable.MakeInvisible()

    # Step 2: YCB visual을 proxy의 자식으로 로드 (physics 없음!)
    ycb_prim_path = f"{proxy_prim_path}/YCB_Visual"
    add_reference_to_stage(usd_path=usd_path, prim_path=ycb_prim_path)
    # NO physics applied to YCB - visual only
"""

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_heading('왜 작동하는가?', level=3)

    reasons = [
        "RMPFlow는 복잡한 mesh collision을 보수적으로 회피",
        "단순한 box collision (DynamicCuboid)은 정확한 접근 허용",
        "Collision Proxy 패턴으로 시각적 표현과 물리 충돌을 분리",
        "YCB visual은 physics 없이 proxy를 따라 이동",
    ]

    for r in reasons:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_page_break()

    # =========================================================================
    # 7. API REFERENCE
    # =========================================================================

    doc.add_heading('7. API 레퍼런스 (API Reference)', level=1)

    doc.add_heading('7.1 주요 Import', level=2)

    imports = """
# Isaac Sim Core API
from isaacsim.core.api import World
from isaacsim.core.api.objects import FixedCuboid, DynamicCuboid
from isaacsim.core.api.objects.ground_plane import GroundPlane

# Robot Manipulators
from isaacsim.robot.manipulators.examples.franka import Franka
from isaacsim.robot.manipulators.examples.franka.controllers import PickPlaceController

# Prims
from isaacsim.core.prims import SingleXFormPrim, SingleRigidPrim

# USD
from pxr import UsdPhysics, UsdGeom, PhysxSchema, Gf

# Utilities
from isaacsim.core.utils.stage import add_reference_to_stage
from isaacsim.storage.native import get_assets_root_path
"""

    p = doc.add_paragraph()
    run = p.add_run(imports)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('7.2 VerificationTracker 클래스', level=2)

    doc.add_paragraph(
        "로봇 동작의 실제 성공/실패를 추적하고 검증하는 커스텀 클래스입니다."
    )

    methods = [
        ("initialize()", "World reset 후 호출 - 객체 초기 위치 저장"),
        ("get_ee_position()", "End Effector 현재 위치 반환"),
        ("get_object_position()", "객체 현재 위치 반환"),
        ("log_phase_status(phase, step)", "Phase별 상세 로깅"),
        ("verify_final_result()", "최종 결과 검증 - 객체가 컨테이너로 이동했는지 확인"),
    ]

    table = doc.add_table(rows=len(methods)+1, cols=2)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = "메서드"
    table.rows[0].cells[1].text = "설명"
    set_cell_shading(table.rows[0].cells[0], 'E5E7EB')
    set_cell_shading(table.rows[0].cells[1], 'E5E7EB')

    for i, (method, desc) in enumerate(methods, 1):
        table.rows[i].cells[0].text = method
        table.rows[i].cells[0].paragraphs[0].runs[0].font.name = 'Consolas'
        table.rows[i].cells[1].text = desc

    doc.add_paragraph()

    doc.add_heading('7.3 RMPFlowObstacleManager 클래스', level=2)

    doc.add_paragraph(
        "RMPFlow 장애물을 동적으로 제어하는 클래스입니다. "
        "Pick 동작 중에는 장애물 회피를 비활성화하고, Place 동작 중에는 활성화합니다."
    )

    methods = [
        ("register_obstacle(obstacle)", "RMPFlow에 장애물 등록"),
        ("disable_for_pick()", "Pick 동작 중 장애물 회피 비활성화"),
        ("enable_after_grasp()", "Grasp 후 장애물 회피 재활성화"),
        ("reset()", "상태 초기화"),
    ]

    table = doc.add_table(rows=len(methods)+1, cols=2)
    table.style = 'Table Grid'

    table.rows[0].cells[0].text = "메서드"
    table.rows[0].cells[1].text = "설명"
    set_cell_shading(table.rows[0].cells[0], 'E5E7EB')
    set_cell_shading(table.rows[0].cells[1], 'E5E7EB')

    for i, (method, desc) in enumerate(methods, 1):
        table.rows[i].cells[0].text = method
        table.rows[i].cells[0].paragraphs[0].runs[0].font.name = 'Consolas'
        table.rows[i].cells[1].text = desc

    doc.add_heading('7.4 YCB_OBJECT_CONFIGS', level=2)

    doc.add_paragraph(
        "지원되는 YCB 객체 설정 딕셔너리입니다. "
        "Franka 그리퍼 최대 열림(80mm) 이하의 객체만 선택했습니다."
    )

    ycb_objects = [
        ("010_potted_meat_can", "Potted Meat Can", "65-70mm", "0.35kg"),
        ("007_tuna_fish_can", "Tuna Fish Can", "~35mm", "0.20kg"),
        ("061_foam_brick", "Foam Brick", "38-50mm", "0.05kg"),
    ]

    table = doc.add_table(rows=len(ycb_objects)+1, cols=4)
    table.style = 'Table Grid'

    headers = ["Object ID", "Display Name", "Width", "Mass"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'E5E7EB')

    for i, (oid, name, width, mass) in enumerate(ycb_objects, 1):
        table.rows[i].cells[0].text = oid
        table.rows[i].cells[1].text = name
        table.rows[i].cells[2].text = width
        table.rows[i].cells[3].text = mass

    doc.add_page_break()

    # =========================================================================
    # 8. TEST RESULTS
    # =========================================================================

    doc.add_heading('8. 테스트 결과 (Test Results)', level=1)

    doc.add_heading('8.1 테스트 이력', level=2)

    doc.add_paragraph(
        "총 26회의 테스트를 통해 문제를 발견하고 해결했습니다."
    )

    test_history = [
        ("1-6회", "YCB 바나나, 테이블 높이 조정", "실패"),
        ("7-16회", "RMPFlow 장애물 문제 확인", "실패"),
        ("17-18회", "RMPFlow disable_obstacle 시도", "효과 없음"),
        ("19회", "테이블 없이 테스트", "EE XY 정확 (0.0147m)"),
        ("20-24회", "페데스탈 설계", "큐브가 밀려남"),
        ("25회", "VisualCuboid 테스트", "EE 정확도 확인"),
        ("26회", "그리퍼 초기화 추가", "성공! (7.1mm)"),
    ]

    table = doc.add_table(rows=len(test_history)+1, cols=3)
    table.style = 'Table Grid'

    headers = ["테스트", "내용", "결과"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], '4F46E5')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for i, (test, content, result) in enumerate(test_history, 1):
        table.rows[i].cells[0].text = test
        table.rows[i].cells[1].text = content
        table.rows[i].cells[2].text = result

    doc.add_paragraph()

    doc.add_heading('8.2 최종 결과', level=2)

    final_results = [
        ("DynamicCuboid", "0.093m", "7.1mm", "SUCCESS"),
        ("Potted Meat Can (YCB)", "0.096m", "11.1mm", "SUCCESS"),
        ("Tuna Fish Can (YCB)", "0.096m", "11.1mm", "SUCCESS"),
        ("Foam Brick (YCB)", "0.096m", "11.1mm", "SUCCESS"),
    ]

    table = doc.add_table(rows=len(final_results)+1, cols=4)
    table.style = 'Table Grid'

    headers = ["객체", "Phase 4 EE Z", "최종 정확도", "결과"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], '10B981')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for i, (obj, ee_z, acc, result) in enumerate(final_results, 1):
        table.rows[i].cells[0].text = obj
        table.rows[i].cells[1].text = ee_z
        table.rows[i].cells[2].text = acc
        table.rows[i].cells[3].text = result

    doc.add_page_break()

    # =========================================================================
    # 9. HOW TO RUN
    # =========================================================================

    doc.add_heading('9. 실행 방법 (How to Run)', level=1)

    doc.add_heading('9.1 GUI 모드', level=2)

    code = """
cd ~/isaac_sim_4.5
./python.sh ~/ClaudeCode_PlanMode_PickAndPlace/franka_ycb_pick_place.py

# GUI 창이 열리면 'S' 키를 눌러 시작
"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_heading('키보드 컨트롤', level=3)

    keys = [
        ("S", "시뮬레이션 시작"),
        ("R", "리셋"),
        ("1", "Potted Meat Can 선택"),
        ("2", "Tuna Fish Can 선택"),
        ("3", "Foam Brick 선택"),
    ]

    table = doc.add_table(rows=len(keys), cols=2)
    table.style = 'Table Grid'

    for i, (key, desc) in enumerate(keys):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = desc

    doc.add_paragraph()

    doc.add_heading('9.2 Headless 모드', level=2)

    doc.add_paragraph(
        "GUI 없이 자동으로 테스트를 실행합니다. CI/CD 또는 배치 테스트에 유용합니다."
    )

    code = """
# franka_ycb_pick_place.py에서 headless 설정 변경
CONFIG = {
    "headless": True,  # GUI 없이 실행
    ...
}

# 실행
cd ~/isaac_sim_4.5
./python.sh ~/ClaudeCode_PlanMode_PickAndPlace/franka_ycb_pick_place.py
"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_heading('9.3 객체 변경', level=2)

    code = """
# franka_ycb_pick_place.py에서 설정 변경

# DynamicCuboid 사용 (기본, 가장 안정적)
USE_YCB_OBJECTS = False

# YCB 객체 사용
USE_YCB_OBJECTS = True
SELECTED_YCB_OBJECT = "010_potted_meat_can"  # 또는 다른 YCB 객체
"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_page_break()

    # =========================================================================
    # 10. FUTURE IMPROVEMENTS
    # =========================================================================

    doc.add_heading('10. 향후 개선 사항 (Future Improvements)', level=1)

    improvements = [
        ("테이블 지원 추가",
         "현재는 테이블 없이 GroundPlane만 사용합니다. RMPFlow 설정을 최적화하여 "
         "테이블이 있어도 정상 동작하도록 개선할 수 있습니다."),

        ("더 많은 YCB 객체 지원",
         "현재 3개의 YCB 객체만 지원합니다. Collision Proxy 패턴을 적용하여 "
         "더 많은 YCB 객체를 추가할 수 있습니다."),

        ("Task 클래스 기반 리팩토링",
         "Isaac Sim의 BaseTask 클래스를 상속하여 코드를 더 구조화할 수 있습니다."),

        ("센서 기반 피드백",
         "현재는 위치 기반 제어만 사용합니다. Force/Torque 센서를 추가하여 "
         "더 정밀한 그립 제어가 가능합니다."),

        ("다중 객체 처리",
         "현재는 한 번에 하나의 객체만 처리합니다. 여러 객체를 순차적으로 "
         "또는 병렬로 처리하는 기능을 추가할 수 있습니다."),
    ]

    for title, desc in improvements:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}")
        run.bold = True
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =========================================================================
    # 11. REFERENCES
    # =========================================================================

    doc.add_heading('11. 참고 자료 (References)', level=1)

    doc.add_heading('11.1 프로젝트 파일', level=2)

    refs = [
        ("메인 스크립트", "franka_ycb_pick_place.py"),
        ("테스트 스크립트", "test_headless.py"),
        ("작업 로그", "work_log.txt"),
        ("3D 프레젠테이션", "project_presentation_v2.html"),
        ("GitHub", "https://github.com/scholarchoi-yjchoi/ClaudeCode_PlanMode_PickAndPlace"),
    ]

    for name, path in refs:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run = p.add_run(path)
        run.font.name = 'Consolas'

    doc.add_heading('11.2 Isaac Sim 공식 자료', level=2)

    isaac_refs = [
        ("공식 예제", "~/isaac_sim_4.5/standalone_examples/api/isaacsim.robot.manipulators/franka_pick_up.py"),
        ("Franka API", "isaacsim.robot.manipulators.examples.franka"),
        ("Core API", "isaacsim.core.api"),
        ("YCB Assets", "Isaac/Props/YCB/Axis_Aligned/"),
    ]

    for name, path in isaac_refs:
        p = doc.add_paragraph()
        run = p.add_run(f"{name}: ")
        run.bold = True
        run = p.add_run(path)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    doc.add_heading('11.3 외부 자료', level=2)

    external = [
        "NVIDIA Isaac Sim Documentation: https://docs.isaacsim.omniverse.nvidia.com/",
        "YCB Dataset: https://www.ycbbenchmarks.com/",
        "Franka Emika Panda: https://www.franka.de/",
        "RMPFlow Paper: https://arxiv.org/abs/2103.04988",
    ]

    for ref in external:
        doc.add_paragraph(ref, style='List Bullet')

    # =========================================================================
    # SAVE DOCUMENT
    # =========================================================================

    output_path = "/home/yjchoi/ClaudeCode_PlanMode_PickAndPlace/Isaac_Sim_Pick_and_Place_Technical_Report_v2.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_document()
